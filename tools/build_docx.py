from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def add_table(doc: Document, rows: list[str]) -> None:
    parsed = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    if len(parsed) < 2:
        return
    if all(re.fullmatch(r":?-{3,}:?", cell) for cell in parsed[1]):
        parsed.pop(1)
    cols = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = "Table Grid"
    for row_index, row in enumerate(parsed):
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.text = row[col_index] if col_index < len(row) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True
    doc.add_paragraph()


def build(input_path: Path, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    document.styles["Title"].font.name = "黑体"
    document.styles["Title"].font.size = Pt(20)
    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = document.styles[style_name]
        style.font.name = "黑体"
        style.font.size = Pt(size)

    lines = input_path.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        if line.strip().startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.25)
                for code_line in code_lines:
                    run = paragraph.add_run(code_line + "\n")
                    run.font.name = "Consolas"
                    run.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|"):
            table_rows = []
            while index < len(lines) and lines[index].startswith("|"):
                table_rows.append(lines[index])
                index += 1
            add_table(document, table_rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1 and not document.paragraphs:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(paragraph, text)
            else:
                paragraph = document.add_paragraph(style=f"Heading {level}")
                add_inline(paragraph, text)
            index += 1
            continue
        if line.strip() == "---":
            document.add_paragraph("-" * 60)
            index += 1
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(line[2:])
            run.italic = True
            index += 1
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        add_inline(paragraph, line)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_docx.py INPUT.md OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
